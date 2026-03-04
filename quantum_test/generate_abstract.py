from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page margins (1 inch all around, standard AMIA) ──
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.0

# ── Title ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MedGenie: Privacy-Preserving Clinical Text Annotation Platform\nwith LLM Integration')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

# ── Authors ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Xiangyu Zeng, MS')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
sup = p.add_run('1')
sup.font.size = Pt(7)
sup.font.name = 'Times New Roman'
sup.font.superscript = True
run = p.add_run(', Yanshan Wang, PhD')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
sup = p.add_run('1')
sup.font.size = Pt(7)
sup.font.name = 'Times New Roman'
sup.font.superscript = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sup2 = p2.add_run('1')
sup2.font.size = Pt(7)
sup2.font.name = 'Times New Roman'
sup2.font.superscript = True
run2 = p2.add_run('PittNAIL Lab, Department of Health Information Management, University of Pittsburgh, Pittsburgh, PA, USA')
run2.font.size = Pt(9)
run2.font.name = 'Times New Roman'
run2.italic = True

# ── Helper: add section with bold heading + body ──
def add_section(heading, body_parts):
    """body_parts: list of (text, bold, italic) tuples"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(heading)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    for text, bold, italic in body_parts:
        run = p2.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
    return p2

# ── Background (compressed) ──
add_section('Background', [
    ('Clinical text annotation is foundational to clinical NLP research. Existing open-source tools such as brat and CLAMP often lack integrated pre-annotation capabilities, while commercial platforms (e.g., Label Studio, Generative AI Lab) that incorporate LLMs raise concerns regarding data privacy and regulatory compliance when clinical text containing protected health information (PHI) is processed through cloud-based APIs. We present MedGenie, a modular annotation platform built upon the open-source tool MedTator, combining a clinician-friendly interface with privacy-preserving local LLM integration for secure pre-annotation.', False, False),
])

# ── Methods ──
add_section('Methods', [
    ('MedGenie was built using React, TypeScript, and CodeMirror 6, based upon the open-source tool MedTator. For LLM-assisted pre-annotation, we implemented a three-stage local pipeline (Figure 1): ', False, False),
    ('(1) Semantic Recognition', True, False),
    (' \u2014 a local LLM via Ollama returns keyword\u2013tag pairs rather than character offsets; ', False, False),
    ('(2) Span Localization', True, False),
    (' \u2014 regex matching deterministically resolves precise character positions; ', False, False),
    ('(3) Negation Filtering', True, False),
    (' \u2014 a bidirectional context-window heuristic suppresses annotations in negated contexts. All processing runs entirely on-device.', False, False),
])

# ── Figure 1: Pipeline diagram as a table ──
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)

table = doc.add_table(rows=3, cols=7, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

# Row 0: boxes and arrows
boxes = ['Clinical\nText +\nSchema', '\u2192', 'Local LLM\n(Ollama)\non-device', '\u2192', 'Regex Span\nLocalization\n(deterministic)', '\u2192', 'Negation\nFilter\n(60/30-char)']
labels = ['', '', 'semantic only\n{keyword, tag}', '', '\\b...\\b \u2192 [s,e]\nexact offset', '', 'bidirectional\ncontext window']

for i, text in enumerate(boxes):
    cell = table.cell(0, i)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(7)
    run.font.name = 'Consolas'
    if i in (1, 3, 5):  # arrow cells
        run.font.size = Pt(12)

# Row 1: labels
for i, text in enumerate(labels):
    cell = table.cell(1, i)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(6.5)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(100, 100, 100)

# Row 2: merge all for output arrow
merged = table.cell(2, 0).merge(table.cell(2, 6))
p = merged.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('\u2192  Pre-annotations loaded into Annotation Editor')
run.font.size = Pt(7)
run.font.name = 'Consolas'
run.bold = True

# Remove all borders from arrow columns
for row in table.rows:
    for i in (1, 3, 5):
        cell = row.cells[i]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = qn('w:tcBorders')
        existing = tcPr.find(borders)
        if existing is not None:
            tcPr.remove(existing)
        new_borders = '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tcBorders>'
        from lxml import etree
        tcPr.append(etree.fromstring(new_borders))

# Caption
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
run = p.add_run('Figure 1. ')
run.bold = True
run.font.size = Pt(9)
run.font.name = 'Times New Roman'
run = p.add_run('Local LLM annotation pipeline. The model provides only keyword\u2013tag pairs; span positions and negation filtering are resolved deterministically on-device.')
run.font.size = Pt(9)
run.font.name = 'Times New Roman'

# ── Placeholder for screenshots ──
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[Figure 2. Screenshot placeholders \u2014 insert app screenshots here]')
run.font.size = Pt(9)
run.font.name = 'Times New Roman'
run.italic = True
run.font.color.rgb = RGBColor(150, 150, 150)

# ── Results (qualitative, concise) ──
add_section('Results', [
    ('MedGenie reproduces all core annotation features of MedTator, including entity/relation annotation, schema editing, IAA with Cohen\u2019s Kappa, multi-format export, and annotation visualization. ', False, False),
    ('We validated the LLM pipeline on 20 COVID-19 adverse event reports from VAERS against a 17-tag gold standard using qwen3:8b via Ollama. Key findings: (1) per-tag semantic descriptions in the prompt substantially improved recall for non-intuitive mappings (e.g., ', False, False),
    ('"swollen" \u2192 Myalgia', False, True),
    ('); (2) short-keyword constraint with overlap deduplication effectively reduced false positives; (3) negation filtering suppressed spurious annotations (e.g., ', False, False),
    ('"denies chest pain"', False, True),
    (') while preserving true positives in affirmative clauses. These findings directly shaped the final pipeline design.', False, False),
])

# ── Conclusion ──
add_section('Conclusion', [
    ('By decoupling keyword extraction from span resolution and applying rule-based negation filtering, MedGenie provides LLM-assisted pre-annotation while keeping all clinical data on-device. The platform is being prepared for Electron-based desktop packaging to enable fully offline clinical annotation workflows. Source code is publicly available at https://github.com/PittNAIL/MedTator.git.', False, False),
])

# ── Save ──
output_path = r'C:\Users\del\Desktop\Work\MedTator\MedTator_AMIA2026_Abstract.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
